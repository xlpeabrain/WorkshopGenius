"""Extract plain text from uploaded PDF and DOCX files."""

import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    result = "\n\n".join(pages)
    if not result.strip():
        raise ValueError(
            "Could not extract any text from this PDF. "
            "It may be a scanned image. Please paste the text manually."
        )
    return result


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    result = "\n".join(parts)
    if not result.strip():
        raise ValueError("Could not extract any text from this DOCX file.")
    return result


def extract_text_from_upload(filename: str, file_bytes: bytes) -> str:
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif name_lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(
            f"Unsupported file type: {filename}. Please upload a PDF, DOCX, or TXT file."
        )
